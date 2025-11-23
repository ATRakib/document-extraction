# app/utils/document_processor.py
import os
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
from pathlib import Path

class DocumentProcessor:
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.tiff'}
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_image(file_path: str) -> str:
        text = ""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
        except Exception as e:
            raise Exception(f"Error extracting Image: {str(e)}")
        return text

    @classmethod
    def extract_text(cls, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return cls.extract_text_from_docx(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff']:
            return cls.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")